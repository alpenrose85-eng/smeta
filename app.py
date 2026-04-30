import io
import json
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document

APP_DIR = Path(__file__).resolve().parent
DATA_PATH = APP_DIR / "data" / "tests.json"
DEFAULT_SERVICES = [
    {
        "code": "chem_perlitic",
        "name": "Химический анализ перлитной трубы (12Х1МФ)",
        "price": 20000,
        "unit": "за образец",
    },
    {
        "code": "chem_austenitic",
        "name": "Химический анализ аустенитной трубы (12Х18Н12Т)",
        "price": 20000,
        "unit": "за образец",
    },
    {
        "code": "chem_steel20",
        "name": "Химический анализ углеродистой трубы (сталь 20)",
        "price": 10000,
        "unit": "за образец",
    },
    {
        "code": "mech_room",
        "name": "Механические испытания на кратковременный разрыв (3 образца с одной трубы) при комнатной температуре",
        "price": 4000,
        "unit": "за образец",
    },
    {
        "code": "mech_hot",
        "name": "Механические испытания на кратковременный разрыв (3 образца с одной трубы) при повышенной температуре",
        "price": 7000,
        "unit": "за образец",
    },
    {
        "code": "sample_micro",
        "name": "Изготовление образцов для механических испытаний (микро)",
        "price": 2000,
        "unit": "за изготовление",
    },
    {
        "code": "sample_classic",
        "name": "Изготовление образцов для механических испытаний (классические)",
        "price": 3000,
        "unit": "за изготовление",
    },
    {
        "code": "sample_long_strength",
        "name": "Изготовление образца на длительную прочность",
        "price": 3000,
        "unit": "за образец",
    },
    {
        "code": "slif_prep",
        "name": "Изготовление шлифа и его подготовка",
        "price": 10000,
        "unit": "за образец",
    },
    {
        "code": "slif_measurement",
        "name": "Измерение фактических размеров",
        "price": 1000,
        "unit": "за шлиф",
    },
    {
        "code": "carbide_analysis",
        "name": "Карбидный анализ",
        "price": 20000,
        "unit": "за образец",
    },
    {
        "code": "electron_microscopy",
        "name": "Электронная микроскопия",
        "price": 20000,
        "unit": "за образец",
    },
    {
        "code": "long_strength",
        "name": "Длительная прочность",
        "price": 200,
        "unit": "за час",
    },
    {
        "code": "damage_investigation",
        "name": "Расследование причины повреждения (анализ)",
        "price": 100000,
        "unit": "за услугу",
    },
    {
        "code": "long_strength_report",
        "name": "Анализ и подготовка заключения о длительной прочности",
        "price": 20000,
        "unit": "за образец",
    },
    {
        "code": "resource_report",
        "name": "Подготовка заключения об оценке остаточного ресурса и расчет",
        "price": 50000,
        "unit": "за поверхность нагрева",
    },
]
STEEL_OPTIONS = ["Сталь 20", "12Х1МФ", "12Х18Н12Т"]
STEEL_SERVICE_RULES = {
    "Сталь 20": [
        ("chem_steel20", 1),
        ("slif_prep", 1),
    ],
    "12Х1МФ": [
        ("chem_perlitic", 1),
        ("carbide_analysis", 1),
        ("slif_prep", 1),
    ],
    "12Х18Н12Т": [
        ("chem_austenitic", 1),
        ("slif_prep", 1),
    ],
}
STEEL20_MECH_RULES = [
    ("mech_room", 1),
    ("mech_hot", 1),
    ("sample_micro", 6),
]
SYSTEM_CODES = {"resource_report", "damage_investigation"}
RESULT_COLUMNS = ["Источник", "Услуга", "Цена", "Количество", "Единица", "Сумма"]


def merge_with_default_services(services: list[dict]) -> list[dict]:
    existing_codes = {service["code"] for service in services}
    merged_services = [service.copy() for service in services]
    for default_service in DEFAULT_SERVICES:
        if default_service["code"] not in existing_codes:
            merged_services.append(default_service.copy())
    return merged_services


@st.cache_data
def load_services() -> list[dict]:
    if not DATA_PATH.exists():
        return [service.copy() for service in DEFAULT_SERVICES]
    with DATA_PATH.open(encoding="utf-8") as file:
        services = json.load(file)
    return merge_with_default_services(services)


def save_services(services: list[dict]) -> None:
    DATA_PATH.parent.mkdir(parents=True, exist_ok=True)
    with DATA_PATH.open("w", encoding="utf-8") as file:
        json.dump(services, file, ensure_ascii=False, indent=2)
    load_services.clear()


def services_map(services: list[dict]) -> dict[str, dict]:
    return {service["code"]: service for service in services}


def money(value: float) -> str:
    return f"{int(value):,} ₽".replace(",", " ")


def rows_to_dataframe(rows: list[dict]) -> pd.DataFrame:
    if not rows:
        return pd.DataFrame(columns=RESULT_COLUMNS)
    return pd.DataFrame(rows, columns=RESULT_COLUMNS)


def add_row(rows: list[dict], source: str, service_name: str, price: float, quantity: float, unit: str) -> None:
    rows.append(
        {
            "Источник": source,
            "Услуга": service_name,
            "Цена": price,
            "Количество": quantity,
            "Единица": unit,
            "Сумма": price * quantity,
        }
    )


def add_slif_measurement_row(rows: list[dict], service_by_code: dict[str, dict], quantity: float, source: str) -> None:
    measurement_service = service_by_code.get("slif_measurement")
    if measurement_service and quantity > 0:
        add_row(
            rows,
            source=source,
            service_name=measurement_service["name"],
            price=measurement_service["price"],
            quantity=quantity,
            unit=measurement_service["unit"],
        )


def calc_steel_rows(
    service_by_code: dict[str, dict],
    steel_counts: dict[str, int],
    steel20_mech_count: int,
) -> list[dict]:
    rows: list[dict] = []
    for steel_name, sample_count in steel_counts.items():
        if sample_count <= 0:
            continue
        for code, multiplier in STEEL_SERVICE_RULES[steel_name]:
            service = service_by_code.get(code)
            if not service:
                continue
            quantity = sample_count * multiplier
            add_row(
                rows,
                source=f"Авто: {steel_name}",
                service_name=service["name"],
                price=service["price"],
                quantity=quantity,
                unit=service["unit"],
            )
            if code == "slif_prep":
                add_slif_measurement_row(rows, service_by_code, quantity, f"Авто: {steel_name}")

    if steel20_mech_count > 0:
        for code, multiplier in STEEL20_MECH_RULES:
            service = service_by_code.get(code)
            if not service:
                continue
            add_row(
                rows,
                source="Авто: Сталь 20 (механические испытания)",
                service_name=service["name"],
                price=service["price"],
                quantity=steel20_mech_count * multiplier,
                unit=service["unit"],
            )
    return rows


def calc_extra_rows(service_by_code: dict[str, dict], selected_codes: list[str]) -> list[dict]:
    rows: list[dict] = []
    for code in selected_codes:
        service = service_by_code[code]
        st.markdown(f"**{service['name']}**")

        if code == "long_strength":
            col1, col2 = st.columns(2)
            sample_count = col1.number_input(
                "Количество образцов",
                min_value=0,
                value=0,
                step=1,
                key=f"{code}_samples",
            )
            hours = col2.number_input(
                "База испытаний, часов на образец",
                min_value=0,
                value=0,
                step=1,
                key=f"{code}_hours",
            )
            if sample_count > 0:
                long_sample_service = service_by_code.get("sample_long_strength")
                if long_sample_service:
                    add_row(
                        rows,
                        source="Авто для длительной прочности",
                        service_name=long_sample_service["name"],
                        price=long_sample_service["price"],
                        quantity=sample_count,
                        unit=long_sample_service["unit"],
                    )

                report_service = service_by_code.get("long_strength_report")
                if report_service:
                    add_row(
                        rows,
                        source="Авто для длительной прочности",
                        service_name=report_service["name"],
                        price=report_service["price"],
                        quantity=sample_count,
                        unit=report_service["unit"],
                    )
            if sample_count > 0 and hours > 0:
                add_row(
                    rows,
                    source="Дополнительная услуга",
                    service_name=service["name"],
                    price=service["price"],
                    quantity=sample_count * hours,
                    unit="час",
                )
        else:
            quantity = st.number_input(
                "Количество",
                min_value=0,
                value=0,
                step=1,
                key=f"{code}_qty",
            )
            if quantity > 0:
                add_row(
                    rows,
                    source="Дополнительная услуга",
                    service_name=service["name"],
                    price=service["price"],
                    quantity=quantity,
                    unit=service["unit"],
                )
                if code == "slif_prep":
                    add_slif_measurement_row(rows, service_by_code, quantity, "Авто для шлифа")
    return rows


def to_csv_bytes(df: pd.DataFrame) -> bytes:
    buffer = io.StringIO()
    export_df = df.copy()
    if not export_df.empty:
        export_df["Цена"] = export_df["Цена"].astype(int)
        export_df["Сумма"] = export_df["Сумма"].astype(int)
    export_df.to_csv(buffer, index=False)
    return buffer.getvalue().encode("utf-8-sig")


def to_docx_bytes(calc_type: str, total: float, df: pd.DataFrame) -> bytes:
    document = Document()
    document.add_heading("Смета", level=1)
    document.add_paragraph(f"Тип расчета: {calc_type}")
    document.add_paragraph(f"Итоговая стоимость: {money(total)}")

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["№", "Услуга", "Цена", "Количество", "Сумма"]
    for idx, title in enumerate(headers):
        table.rows[0].cells[idx].text = title

    for number, (_, row) in enumerate(df.iterrows(), start=1):
        cells = table.add_row().cells
        cells[0].text = str(number)
        cells[1].text = str(row["Услуга"])
        cells[2].text = money(row["Цена"])
        cells[3].text = f"{row['Количество']} {row['Единица']}"
        cells[4].text = money(row["Сумма"])

    document.add_paragraph("")
    document.add_paragraph(f"Итого: {money(total)}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def render_services_tab(services: list[dict]) -> None:
    st.subheader("Справочник услуг")
    st.caption("Здесь можно изменить стоимость стандартных услуг и добавить новые услуги для будущих расчетов.")

    editable_df = pd.DataFrame(services)
    editable_df = editable_df[["code", "name", "price", "unit"]]
    editable_df.columns = ["Код", "Название услуги", "Стоимость", "Единица"]

    edited_df = st.data_editor(
        editable_df,
        use_container_width=True,
        hide_index=True,
        disabled=["Код"],
        column_config={
            "Стоимость": st.column_config.NumberColumn("Стоимость", min_value=0, step=100),
        },
        key="services_editor",
    )

    col_save, col_info = st.columns([1, 2])
    if col_save.button("Сохранить изменения в справочнике", type="primary"):
        updated_services = []
        for original, (_, row) in zip(services, edited_df.iterrows()):
            updated_services.append(
                {
                    "code": original["code"],
                    "name": str(row["Название услуги"]).strip(),
                    "price": int(row["Стоимость"]),
                    "unit": str(row["Единица"]).strip(),
                }
            )
        save_services(updated_services)
        st.success("Изменения сохранены.")
        st.rerun()
    col_info.info("Коды служебные. Их лучше не менять, поэтому они заблокированы.")

    st.divider()
    st.subheader("Добавить новую услугу")
    with st.form("add_service_form"):
        name = st.text_input("Название услуги")
        price = st.number_input("Стоимость", min_value=0, value=0, step=100)
        unit = st.text_input("Единица", value="за образец")
        submitted = st.form_submit_button("Добавить услугу")

    if submitted:
        cleaned_name = name.strip()
        cleaned_unit = unit.strip() or "за образец"
        if not cleaned_name:
            st.warning("Введите название услуги.")
        else:
            existing_codes = {service["code"] for service in services}
            base_code = "custom_" + "_".join(cleaned_name.lower().split())
            new_code = base_code
            counter = 2
            while new_code in existing_codes:
                new_code = f"{base_code}_{counter}"
                counter += 1

            updated_services = services + [
                {
                    "code": new_code,
                    "name": cleaned_name,
                    "price": int(price),
                    "unit": cleaned_unit,
                }
            ]
            save_services(updated_services)
            st.success(f"Услуга «{cleaned_name}» добавлена.")
            st.rerun()


def render_calculation_tab(services: list[dict]) -> None:
    st.subheader("Расчет")
    service_by_code = services_map(services)

    calc_type = st.radio(
        "Что рассчитываем?",
        options=["Оценка остаточного ресурса", "Расследование причины повреждения"],
        horizontal=True,
    )

    rows: list[dict] = []

    if calc_type == "Оценка остаточного ресурса":
        surface_count = st.number_input(
            "Количество поверхностей нагрева",
            min_value=0,
            value=0,
            step=1,
        )
        if surface_count > 0 and "resource_report" in service_by_code:
            service = service_by_code["resource_report"]
            add_row(
                rows,
                source="Базовая услуга",
                service_name=service["name"],
                price=service["price"],
                quantity=surface_count,
                unit=service["unit"],
            )
    else:
        if "damage_investigation" in service_by_code:
            service = service_by_code["damage_investigation"]
            add_row(
                rows,
                source="Базовая услуга",
                service_name=service["name"],
                price=service["price"],
                quantity=1,
                unit=service["unit"],
            )

    st.divider()
    st.markdown("### Количество образцов по маркам стали")
    st.caption("Заполните количество образцов напротив каждой марки стали.")

    steel_counts: dict[str, int] = {}

    steel20_count = st.number_input(
        "Сталь 20",
        min_value=0,
        value=0,
        step=1,
        key="steel_Сталь 20",
    )
    steel_counts["Сталь 20"] = steel20_count

    steel20_mech_enabled = st.checkbox(
        "Механические испытания",
        key="steel20_mech_enabled",
    )
    prev_enabled = st.session_state.get("steel20_mech_enabled_prev", False)
    if steel20_mech_enabled:
        if not prev_enabled or "steel20_mech_count_input" not in st.session_state:
            st.session_state["steel20_mech_count_input"] = steel20_count
        else:
            st.session_state["steel20_mech_count_input"] = min(
                st.session_state.get("steel20_mech_count_input", steel20_count),
                steel20_count,
            )
    else:
        st.session_state["steel20_mech_count_input"] = 0
    st.session_state["steel20_mech_enabled_prev"] = steel20_mech_enabled

    steel20_mech_count = st.number_input(
        "Количество образцов для механических испытаний",
        min_value=0,
        max_value=steel20_count,
        step=1,
        key="steel20_mech_count_input",
        disabled=not steel20_mech_enabled,
    )

    for steel_name in ["12Х1МФ", "12Х18Н12Т"]:
        steel_counts[steel_name] = st.number_input(
            steel_name,
            min_value=0,
            value=0,
            step=1,
            key=f"steel_{steel_name}",
        )

    rows.extend(calc_steel_rows(service_by_code, steel_counts, steel20_mech_count if steel20_mech_enabled else 0))

    st.divider()
    st.markdown("### Дополнительные услуги")
    additional_service_codes = [code for code in service_by_code if code not in SYSTEM_CODES]
    selected_extra_codes = st.multiselect(
        "Выберите дополнительные услуги",
        options=additional_service_codes,
        format_func=lambda code: service_by_code[code]["name"],
    )

    if selected_extra_codes:
        rows.extend(calc_extra_rows(service_by_code, selected_extra_codes))

    result_df = rows_to_dataframe(rows)
    total = result_df["Сумма"].sum() if not result_df.empty else 0

    st.session_state["estimate_df"] = result_df.copy()
    st.session_state["estimate_total"] = total
    st.session_state["estimate_calc_type"] = calc_type

    st.divider()
    st.markdown("### Детализация расчета")
    if result_df.empty:
        st.info("Пока нет данных для расчета. Заполните параметры выше.")
        return

    display_df = result_df.copy()
    display_df["Цена"] = display_df["Цена"].apply(money)
    display_df["Сумма"] = display_df["Сумма"].apply(money)
    st.dataframe(display_df, use_container_width=True, hide_index=True)
    st.metric("Итоговая стоимость", money(total))

    csv_bytes = to_csv_bytes(result_df)
    st.download_button(
        label="Скачать расчет (CSV)",
        data=csv_bytes,
        file_name="raschet_stoimosti_uslug.csv",
        mime="text/csv",
    )


def render_estimate_tab() -> None:
    st.subheader("Смета")
    st.caption("Готовая смета формируется на основе текущего расчета.")

    result_df = st.session_state.get("estimate_df")
    total = st.session_state.get("estimate_total", 0)
    calc_type = st.session_state.get("estimate_calc_type", "Не выбран")

    if result_df is None or result_df.empty:
        st.info("Сначала заполните вкладку «Расчет». После этого здесь появится смета.")
        return

    estimate_df = result_df[["Услуга", "Цена", "Количество", "Единица", "Сумма"]].copy()
    display_df = estimate_df.copy()
    display_df.insert(0, "№", range(1, len(display_df) + 1))
    display_df["Цена"] = display_df["Цена"].apply(money)
    display_df["Сумма"] = display_df["Сумма"].apply(money)
    display_df["Количество"] = display_df.apply(
        lambda row: f"{row['Количество']} {row['Единица']}", axis=1
    )
    display_df = display_df[["№", "Услуга", "Цена", "Количество", "Сумма"]]

    st.markdown(f"**Тип расчета:** {calc_type}")
    st.dataframe(display_df, use_container_width=True, hide_index=True)
    st.metric("Итого по смете", money(total))

    docx_bytes = to_docx_bytes(calc_type, total, result_df)
    st.download_button(
        label="Скачать смету (DOCX)",
        data=docx_bytes,
        file_name="smeta.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
    )


def main() -> None:
    st.set_page_config(page_title="Расчет стоимости услуг", layout="wide")
    st.title("Расчет стоимости услуг")
    st.caption("Справочник услуг хранится в файле и используется во всех последующих расчетах.")

    services = load_services()
    tabs = st.tabs(["Расчет", "Смета", "Справочник услуг"])

    with tabs[0]:
        render_calculation_tab(services)

    with tabs[1]:
        render_estimate_tab()

    with tabs[2]:
        render_services_tab(services)


if __name__ == "__main__":
    main()
